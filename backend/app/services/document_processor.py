"""Document processing service for PDF/DOCX parsing and semantic chunking.

This module provides the complete document processing pipeline:
- PDF extraction using pymupdf4llm (clean Markdown output)
- DOCX extraction using python-docx
- Semantic chunking with RecursiveCharacterTextSplitter
- Async pipeline for background processing
- Task status tracking for progress visibility

CRITICAL: Uses semantic chunking to prevent Pitfall #2 (poor chunking strategy).
"""

import logging
import os
import uuid
from typing import Dict, List

import pymupdf4llm
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import settings
from app.utils.task_tracker import TaskStatus, task_tracker

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using pymupdf4llm with raw-text and OCR fallbacks.

    Tries in order:
    1. pymupdf4llm markdown extraction (preserves structure)
    2. Raw pymupdf text extraction
    3. OCR via tesseract (for image-based/scanned PDFs)

    Args:
        file_path: Path to the PDF file.

    Returns:
        Extracted text as string.
    """
    import pymupdf

    # 1. Try markdown extraction
    md_text = pymupdf4llm.to_markdown(file_path)
    if md_text.strip():
        return md_text

    # 2. Fallback: raw text extraction
    logger.warning("pymupdf4llm returned empty markdown, trying raw text extraction")
    doc = pymupdf.open(file_path)
    parts = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            parts.append(text)
    doc.close()

    if parts:
        return "\n\n".join(parts)

    # 3. Fallback: OCR via tesseract
    logger.warning("Raw text extraction empty, trying OCR")
    doc = pymupdf.open(file_path)
    ocr_parts = []
    for page in doc:
        tp = page.get_textpage_ocr(tessdata=None, dpi=300, full=True)
        text = page.get_text(textpage=tp)
        if text.strip():
            ocr_parts.append(text)
    doc.close()

    if ocr_parts:
        logger.info(f"OCR extracted text from {len(ocr_parts)} pages")
    return "\n\n".join(ocr_parts)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx.

    Extracts paragraphs and tables, joining with double newlines
    for semantic boundary preservation.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text as string with semantic boundaries.
    """
    doc = Document(file_path)
    parts: List[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract tables
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                table_text.append(row_text)
        if table_text:
            parts.append("\n".join(table_text))

    # Join with double newlines for semantic boundaries
    return "\n\n".join(parts)


def chunk_text(text: str) -> List[Dict]:
    """Chunk text using semantic chunking.

    CRITICAL: Uses RecursiveCharacterTextSplitter, NOT fixed-size chunking.
    Prevents Pitfall #2 (poor chunking strategy) by respecting semantic boundaries.

    The splitter tries separators in order:
    1. Double newlines (paragraph boundaries)
    2. Single newlines (line breaks)
    3. Sentence endings (". ")
    4. Spaces (words)
    5. Empty string (characters)

    Args:
        text: Text to chunk.

    Returns:
        List of chunk dictionaries with 'text' and 'position' keys.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""],  # Respect semantic boundaries
    )

    chunks = splitter.split_text(text)

    return [{"text": chunk, "position": idx} for idx, chunk in enumerate(chunks)]


async def process_document_pipeline(
    file_path: str,
    document_id: str,
    user_id: str,
    filename: str,
    file_size: int = 0,
) -> None:
    """Complete document processing pipeline with status tracking.

    Steps (with status updates):
    1. EXTRACTING: Extract text (PDF/DOCX based on file extension)
    2. CHUNKING: Chunk text (semantic chunking)
    3. EMBEDDING: Generate embeddings for all chunks
    4. SUMMARIZING: Generate document summary using LangChain
    5. INDEXING: Store in Neo4j and Qdrant (with summary)
    6. COMPLETED: Processing finished

    Runs in background task to avoid blocking API (Pitfall #7).

    Args:
        file_path: Path to the uploaded file.
        document_id: UUID for the document.
        user_id: ID of the user who uploaded the document.
        filename: Original filename for metadata.
    """
    # Import here to avoid circular imports
    from app.services.embedding_service import generate_embeddings
    from app.services.indexing_service import (
        store_chunks_in_qdrant,
        store_document_in_neo4j,
    )

    try:
        logger.info(f"Processing document: {filename} (id: {document_id})")

        # Step 1: Extract text
        task_tracker.update(
            document_id, TaskStatus.EXTRACTING, f"Extracting text from {filename}"
        )
        _, ext = os.path.splitext(filename.lower())
        if ext == ".pdf":
            text = extract_text_from_pdf(file_path)
        elif ext in (".docx", ".doc"):
            text = extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        logger.info(f"Extracted {len(text)} characters from {filename}")

        # Step 2: Chunk text
        task_tracker.update(
            document_id, TaskStatus.CHUNKING, "Splitting into semantic chunks"
        )
        chunks = chunk_text(text)
        logger.info(f"Created {len(chunks)} chunks from {filename}")

        if not chunks:
            logger.warning(f"No chunks created from {filename} - empty document?")
            task_tracker.fail(document_id, "Document appears to be empty")
            return

        # Step 3: Generate embeddings
        task_tracker.update(
            document_id,
            TaskStatus.EMBEDDING,
            f"Generating embeddings for {len(chunks)} chunks",
        )
        chunk_texts = [chunk["text"] for chunk in chunks]
        embeddings = await generate_embeddings(chunk_texts)
        logger.info(f"Generated {len(embeddings)} embeddings for {filename}")

        # Prepare chunk data with shared UUIDs for Neo4j/Qdrant linkage
        # CRITICAL: Same ID used in both stores for cross-referencing
        chunk_data = []
        for chunk, embedding in zip(chunks, embeddings):
            chunk_id = str(uuid.uuid4())
            chunk_data.append(
                {
                    "id": chunk_id,
                    "text": chunk["text"],
                    "position": chunk["position"],
                    "vector": embedding,
                    "document_id": document_id,
                    "user_id": user_id,
                }
            )

        # Step 4: Store in databases (summary generated on-demand, not during upload)
        task_tracker.update(document_id, TaskStatus.INDEXING, "Storing in database")
        store_document_in_neo4j(
            document_id=document_id,
            user_id=user_id,
            filename=filename,
            chunks=chunk_data,
            summary="",
            file_size=file_size,
        )
        logger.info(f"Stored document and chunks in Neo4j for {filename}")

        store_chunks_in_qdrant(chunk_data)
        logger.info(f"Stored vectors in Qdrant for {filename}")

        # Step 5b: Extract and store entities (GraphRAG)
        if settings.GRAPHRAG_ENABLED:
            task_tracker.update(
                document_id,
                TaskStatus.EXTRACTING_ENTITIES,
                f"Extracting entities from {len(chunk_data)} chunks",
            )
            from app.services.entity_extraction_service import extract_entities_batch
            from app.services.indexing_service import store_entities_in_neo4j

            entity_results = await extract_entities_batch(chunk_data, document_id=document_id)
            for chunk_datum, entity_result in zip(chunk_data, entity_results):
                if entity_result.get("entities") or entity_result.get("relationships"):
                    store_entities_in_neo4j(
                        chunk_id=chunk_datum["id"],
                        entities=entity_result.get("entities", []),
                        relationships=entity_result.get("relationships", []),
                    )
            logger.info(f"Stored entities in Neo4j for {filename}")

        # Step 6: Complete
        task_tracker.complete(document_id, "Document processed successfully")
        logger.info(f"Successfully processed document: {filename} (id: {document_id})")

    except Exception as e:
        logger.error(f"Error processing document {filename}: {e}")
        task_tracker.fail(document_id, str(e))
        raise

    finally:
        # Clean up temp file
        if os.path.exists(file_path):
            os.unlink(file_path)
            logger.debug(f"Cleaned up temp file: {file_path}")
